"""Multi-format document reader for the Document_QA skill.

Bu modül `.pdf`, `.docx`, `.txt` ve `.md` uzantılı dosyaları tek bir
``read(path) -> str`` arayüzünden okur (Requirement 20.1, Design §8). Çıkış
tek string'tir ve whitespace normalize edilmiştir (Property 19): ardışık
boşluk/sekme/yeni satır karakterleri tek bir boşluğa indirgenir, baş/son
boşluklar atılır.

Tasarım kararları:
    * Uzantıya göre dispatch — okuyucu seçimi case-insensitive.
    * PDF: ``pypdf.PdfReader``; her sayfanın ``extract_text()`` çıktısı
      whitespace ayracıyla birleştirilir.
    * DOCX: ``docx.Document``; tüm paragraflar ve tablo hücreleri birleştirilir
      (gövde metni + tablo metni; başlık/altbilgi taranmaz).
    * TXT / MD: UTF-8 öncelikli okuma; UnicodeDecodeError olursa ``cp1254``
      (Türkçe Windows codepage) ardından latin-1 fallback'i.
    * Bilinmeyen uzantı: ``ValueError`` ile destek dışı olduğu belirtilir.
    * Dosya yoksa ``FileNotFoundError`` (Design §"Error Handling Strategy" —
      "Dosya bulunamadı: {path}").

Bu modül yalnızca metin çıkarmaktan sorumludur; chunking, prompt
oluşturma ve LLM çağrıları sonraki alt görevlerde (11.2, 11.3) eklenir.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Callable

# Whitespace karakterlerinin tamamını (boşluk, tab, yeni satır, NBSP, form
# feed, vertical tab vb. dahil) tek bir ASCII boşluğa indirgeyen regex.
# ``str.split() + " ".join(...)`` yerine regex kullanmamızın sebebi
# ``\u00a0`` (non-breaking space) gibi DOCX/PDF'de sıkça gelen Unicode
# whitespace karakterlerini de yakalamak.
_WHITESPACE_RE = re.compile(r"\s+", flags=re.UNICODE)


def _normalize_whitespace(text: str) -> str:
    """Tüm whitespace karakterlerini tek bir boşluğa indir, uçları kırp."""

    return _WHITESPACE_RE.sub(" ", text).strip()


# ---------------------------------------------------------------------------
# Format-specific readers
# ---------------------------------------------------------------------------


def _read_text(path: Path) -> str:
    """``.txt`` ve ``.md`` dosyaları için düz metin okuma.

    UTF-8 önce denenir; başarısız olursa cp1254 (Türkçe Windows) ve son
    çare olarak latin-1 ile okunur. latin-1 her byte sekansını başarıyla
    çözebildiği için fallback zinciri her zaman tamamlanır.
    """

    for encoding in ("utf-8", "cp1254", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    # latin-1 her zaman başarılı olduğundan teorik olarak buraya düşmeyiz.
    return path.read_text(encoding="latin-1", errors="replace")


def _read_pdf(path: Path) -> str:
    """``pypdf`` ile PDF içeriğini sayfa sayfa çıkar."""

    # Import gecikmeli: pypdf opsiyonel bağımlılık olarak kalır ve test
    # ortamlarında stub'lanabilir.
    from pypdf import PdfReader  # type: ignore[import-untyped]

    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        try:
            extracted = page.extract_text() or ""
        except Exception:  # noqa: BLE001 — PDF parsing hataları yutulur
            extracted = ""
        pages.append(extracted)
    # Sayfaları boşlukla birleştir; whitespace normalize sonradan run-on'u
    # tek boşluğa indireceği için " " kullanmak güvenli.
    return " ".join(pages)


def _read_docx(path: Path) -> str:
    """``python-docx`` ile DOCX gövde metni ve tabloları çıkar."""

    from docx import Document  # type: ignore[import-untyped]

    document = Document(str(path))
    parts: list[str] = []

    # Paragraflar — başlıklar ve gövde metninin tamamı.
    for paragraph in document.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)

    # Tablolar — her hücre kendi başına bir paragraf gibi davranır.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)

    return " ".join(parts)


# ---------------------------------------------------------------------------
# Dispatch
# ---------------------------------------------------------------------------


# Uzantı → reader haritası. Yeni format eklemek isteyen geliştirici buraya
# bir satır ekler ve testlere fixture düşer.
_READERS: dict[str, Callable[[Path], str]] = {
    ".txt": _read_text,
    ".md": _read_text,
    ".pdf": _read_pdf,
    ".docx": _read_docx,
}


def supported_extensions() -> tuple[str, ...]:
    """Desteklenen uzantıların alfabetik tuple'ı (test ve diagnostic için)."""

    return tuple(sorted(_READERS.keys()))


def read(path: str | Path) -> str:
    """Bir belgeyi oku ve whitespace normalize edilmiş tek string döndür.

    Args:
        path: ``.pdf``, ``.docx``, ``.txt`` veya ``.md`` uzantılı bir dosya
            yolu.

    Returns:
        Belgenin tüm metinsel içeriği; ardışık whitespace tek boşluğa
        indirgenmiş ve uçları kırpılmış halde.

    Raises:
        FileNotFoundError: Dosya diskte yoksa.
        ValueError: Uzantı desteklenmiyorsa.
    """

    file_path = Path(path)

    if not file_path.exists():
        # Design § Error Handling: "Dosya bulunamadı: {path}".
        raise FileNotFoundError(f"Dosya bulunamadı: {file_path}")

    if not file_path.is_file():
        raise FileNotFoundError(f"Yol bir dosya değil: {file_path}")

    suffix = file_path.suffix.lower()
    reader = _READERS.get(suffix)
    if reader is None:
        supported = ", ".join(supported_extensions())
        raise ValueError(
            f"Desteklenmeyen belge uzantısı: {suffix!r}. "
            f"Desteklenen: {supported}"
        )

    raw_text = reader(file_path)
    return _normalize_whitespace(raw_text)


__all__ = ["read", "supported_extensions"]
